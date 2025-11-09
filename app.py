#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ZOPTYMALIZOWANY GENERATOR OFERT - LibreOffice + PyMuPDF
- Szybka konwersja DOCX → PDF → JPG
- Pre-rendering JPG szablonów na starcie
- Cache dla produktów
- Jeden plik - wszystko w środku
"""

import os
import json
import tempfile
import shutil
import subprocess
import base64
import hashlib
import threading
from pathlib import Path
from flask import Flask, render_template, request, jsonify, send_file, make_response
from flask_socketio import SocketIO, emit
from datetime import datetime
from docx import Document
import docx.enum.text
from flask_compress import Compress
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    import fitz  # PyMuPDF - super szybkie!
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    print("[WARNING] PyMuPDF nie zainstalowane - używam pdf2image (wolniejsze)")
    from pdf2image import convert_from_path

# ============================================================
# KONFIGURACJA
# ============================================================

app = Flask(__name__)
app.config['SECRET_KEY'] = 'secret_key_for_socketio_12345'
app.config['COMPRESS_MIMETYPES'] = ['application/json', 'text/html', 'text/css', 'application/javascript']
app.config['COMPRESS_LEVEL'] = 6
app.config['COMPRESS_MIN_SIZE'] = 500
Compress(app)

socketio = SocketIO(app, cors_allowed_origins="*")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = os.path.join(BASE_DIR, 'templates')
PRODUKTY_DIR = os.path.join(BASE_DIR, 'produkty')
SAVED_OFFERS_DIR = os.path.join(BASE_DIR, 'saved_offers')
GENERATED_OFFERS_DIR = os.path.join(BASE_DIR, 'generated_offers')
OUT_JPG_DIR = os.path.join(BASE_DIR, 'out_jpg')

# Utwórz foldery
os.makedirs(SAVED_OFFERS_DIR, exist_ok=True)
os.makedirs(GENERATED_OFFERS_DIR, exist_ok=True)
os.makedirs(OUT_JPG_DIR, exist_ok=True)

# Globalne cache
libreoffice_lock = threading.Lock()
conversion_cache = {}  # {file_hash: [list of base64 images]}

# ============================================================
# KONWERSJA DOCX → JPG (LibreOffice + PyMuPDF)
# ============================================================

def find_libreoffice():
    """Znajdź soffice w systemie"""
    paths = [
        '/usr/bin/soffice',
        '/usr/local/bin/soffice',
        'soffice',
        '/Applications/LibreOffice.app/Contents/MacOS/soffice'
    ]
    for path in paths:
        if shutil.which(path) or os.path.exists(path):
            return path
    return None


def docx_to_pdf_libreoffice(docx_path, out_pdf_path):
    """Konwertuj DOCX → PDF używając LibreOffice"""
    soffice = find_libreoffice()
    if not soffice:
        raise RuntimeError("LibreOffice nie znalezione! Zainstaluj: apt install libreoffice")

    outdir = os.path.dirname(out_pdf_path)

    # Mutex - tylko jedna konwersja LibreOffice naraz
    with libreoffice_lock:
        cmd = [
            soffice,
            '--headless',
            '--nologo',
            '--nodefault',
            '--nofirststartwizard',
            '--convert-to', 'pdf:writer_pdf_Export',
            '--outdir', outdir,
            docx_path
        ]

        result = subprocess.run(cmd, capture_output=True, timeout=60)

        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice error: {result.stderr.decode()}")

        # Znajdź wygenerowany PDF
        candidate = Path(outdir) / (Path(docx_path).stem + ".pdf")
        if not candidate.exists():
            pdfs = list(Path(outdir).glob("*.pdf"))
            if not pdfs:
                raise RuntimeError("LibreOffice nie wygenerował PDF")
            candidate = max(pdfs, key=lambda p: p.stat().st_mtime)

        # Przenieś do oczekiwanej lokalizacji
        if str(candidate) != out_pdf_path:
            shutil.move(str(candidate), out_pdf_path)


def pdf_to_jpg_pymupdf(pdf_path, dpi=200, quality=90):
    """Konwertuj PDF → JPG używając PyMuPDF (SUPER FAST!)"""
    if not HAS_PYMUPDF:
        raise RuntimeError("PyMuPDF nie zainstalowane! pip install pymupdf")

    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    images = []

    with fitz.open(pdf_path) as doc:
        for page in doc:
            pix = page.get_pixmap(matrix=mat, alpha=False)

            # Zapisz do bytes
            img_bytes = pix.tobytes("jpeg", jpg_quality=quality)
            img_base64 = base64.b64encode(img_bytes).decode('utf-8')
            images.append(f"data:image/jpeg;base64,{img_base64}")

    return images


def pdf_to_jpg_pdf2image(pdf_path, dpi=200):
    """Fallback: Konwertuj PDF → JPG używając pdf2image"""
    from PIL import Image
    import io

    pages = convert_from_path(pdf_path, dpi=dpi)
    images = []

    for page in pages:
        # Konwertuj do RGB jeśli RGBA
        if page.mode == 'RGBA':
            bg = Image.new('RGB', page.size, (255, 255, 255))
            bg.paste(page, mask=page.split()[3])
            page = bg

        # Zapisz jako JPEG
        buffered = io.BytesIO()
        page.save(buffered, format="JPEG", quality=85, optimize=True)
        img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
        images.append(f"data:image/jpeg;base64,{img_base64}")

    return images


def convert_docx_to_images(docx_path, use_cache=True, progress_callback=None):
    """
    Główna funkcja: DOCX → JPG
    1. Sprawdź cache
    2. DOCX → PDF (LibreOffice)
    3. PDF → JPG (PyMuPDF lub pdf2image)
    """
    # Cache
    if use_cache:
        file_hash = get_file_hash(docx_path)
        if file_hash and file_hash in conversion_cache:
            print(f"[CACHE] ⚡ Hit: {os.path.basename(docx_path)}")
            return conversion_cache[file_hash]

    print(f"[CONVERT] Start: {os.path.basename(docx_path)}")

    if progress_callback:
        progress_callback("Konwersja DOCX → PDF...", 20)

    # Tymczasowy PDF
    with tempfile.TemporaryDirectory(dir=OUT_JPG_DIR) as tmpdir:
        pdf_path = os.path.join(tmpdir, 'out.pdf')

        # DOCX → PDF
        docx_to_pdf_libreoffice(docx_path, pdf_path)

        if progress_callback:
            progress_callback("Konwersja PDF → JPG...", 50)

        # PDF → JPG
        if HAS_PYMUPDF:
            images = pdf_to_jpg_pymupdf(pdf_path, dpi=200, quality=90)
        else:
            images = pdf_to_jpg_pdf2image(pdf_path, dpi=200)

    # Zapisz w cache
    if use_cache:
        file_hash = get_file_hash(docx_path)
        if file_hash:
            conversion_cache[file_hash] = images
            print(f"[CACHE] ✓ Saved: {os.path.basename(docx_path)}")

    print(f"[CONVERT] ✓ Done: {len(images)} stron")
    return images


def get_file_hash(filepath):
    """Hash pliku dla cache"""
    try:
        with open(filepath, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()
    except:
        return None


# ============================================================
# PRE-RENDERING NA STARCIE
# ============================================================

def preload_all_products():
    """Pre-renderuj wszystkie produkty przy starcie"""
    print("\n" + "="*80)
    print("[STARTUP] 🚀 Pre-rendering produktów...")
    print("="*80)

    if not os.path.exists(PRODUKTY_DIR):
        print("[STARTUP] Folder produktów nie istnieje")
        return

    product_files = [f for f in os.listdir(PRODUKTY_DIR) if f.endswith('.docx') and not f.startswith('~$')]
    total = len(product_files)

    print(f"[STARTUP] Znaleziono {total} produktów")

    for idx, filename in enumerate(product_files, 1):
        product_path = os.path.join(PRODUKTY_DIR, filename)
        print(f"[STARTUP] [{idx}/{total}] {filename}...", end=' ')

        try:
            convert_docx_to_images(product_path, use_cache=True)
            print("✓")
        except Exception as e:
            print(f"✗ {e}")

    print(f"[STARTUP] ✅ Cache: {len(conversion_cache)} produktów")
    print("="*80 + "\n")


def preload_templates():
    """Pre-renderuj JPG szablonów WolfTax na starcie"""
    print("\n" + "="*80)
    print("[STARTUP] 🎨 Pre-rendering szablonów WolfTax...")
    print("="*80)

    wolftax_folder = os.path.join(TEMPLATES_DIR, 'wolftax-oferta')
    if not os.path.exists(wolftax_folder):
        print("[STARTUP] Folder wolftax-oferta nie istnieje")
        return

    files = ['Dok1.docx', 'Doc2.docx', 'doc3.docx', 'doc4.docx', 'Dok5.docx', 'Dok6.docx']

    for filename in files:
        filepath = os.path.join(wolftax_folder, filename)
        if not os.path.exists(filepath):
            continue

        print(f"[STARTUP] Renderuję {filename}...", end=' ')

        try:
            # Renderuj do OUT_JPG_DIR jako statyczne JPG
            out_folder = os.path.join(OUT_JPG_DIR, filename.replace('.docx', ''))
            os.makedirs(out_folder, exist_ok=True)

            # Konwertuj
            images = convert_docx_to_images(filepath, use_cache=False)

            # Zapisz JPG na dysk
            for i, img_data in enumerate(images, 1):
                # Wyciągnij base64
                if img_data.startswith('data:image/jpeg;base64,'):
                    img_b64 = img_data.split(',', 1)[1]
                    img_bytes = base64.b64decode(img_b64)

                    jpg_path = os.path.join(out_folder, f'page_{i:04d}.jpg')
                    with open(jpg_path, 'wb') as f:
                        f.write(img_bytes)

            print(f"✓ {len(images)} stron")
        except Exception as e:
            print(f"✗ {e}")

    print("[STARTUP] ✅ Szablony JPG gotowe!")
    print("="*80 + "\n")


def preload_async():
    """Uruchom pre-rendering w tle"""
    thread = threading.Thread(target=lambda: (preload_all_products(), preload_templates()), daemon=True)
    thread.start()


# ============================================================
# POMOCNICZE FUNKCJE
# ============================================================

def send_progress(message, percent):
    """Wyślij progress przez WebSocket"""
    try:
        socketio.emit('conversion_progress', {'message': message, 'percent': percent})
    except:
        pass


def send_page_ready(page_data):
    """Wyślij gotową stronę przez WebSocket"""
    try:
        socketio.emit('page_ready', page_data)
    except:
        pass


def replace_placeholders(doc, data):
    """Zamień {{placeholders}} w dokumencie"""
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key != 'produkty':
                placeholder = '{{' + key + '}}'
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if key != 'produkty':
                        placeholder = '{{' + key + '}}'
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))

    return doc


def add_page_break_to_doc(doc):
    """Dodaj page break na końcu dokumentu"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)
    return doc


def merge_documents(docs):
    """
    Prosta funkcja do łączenia wielu dokumentów DOCX
    Zwraca połączony dokument
    """
    if not docs:
        return Document()

    # Pierwszy dokument jako baza
    merged = docs[0]

    # Dodaj pozostałe dokumenty
    for doc in docs[1:]:
        # Dodaj page break
        merged = add_page_break_to_doc(merged)

        # Skopiuj wszystkie elementy z doc do merged
        for element in doc.element.body:
            merged.element.body.append(element)

    return merged


def generate_table_of_contents(selected_products, product_custom_fields, start_page=5):
    """Generuj spis treści dla WolfTax"""
    toc_lines = []
    current_page = start_page

    for idx, product_id in enumerate(selected_products, 1):
        product_data = product_custom_fields.get(product_id, {})
        title = product_data.get('title') or product_data.get('nazwa') or f'Produkt {product_id}'

        base_text = f"Usługa {idx} – {title}"
        dots_count = max(60 - len(base_text), 10)
        dots = '…' * dots_count

        line = f"§\t{base_text} {dots}  {current_page:02d}"
        toc_lines.append(line)

        # Policz strony produktu
        product_path = os.path.join(PRODUKTY_DIR, f'{product_id}.docx')
        file_hash = get_file_hash(product_path)
        if file_hash and file_hash in conversion_cache:
            product_pages = len(conversion_cache[file_hash])
        else:
            product_pages = 1

        current_page += product_pages

    return '\n'.join(toc_lines)


def inject_toc_into_doc(doc, toc_text):
    """Wstaw spis treści do dokumentu"""
    injected = False

    for para in doc.paragraphs:
        if '{{SPIS_TRESCI}}' in para.text or '{{TOC}}' in para.text:
            para.text = para.text.replace('{{SPIS_TRESCI}}', toc_text)
            para.text = para.text.replace('{{TOC}}', toc_text)
            injected = True
            break

    if not injected:
        doc.add_paragraph(toc_text)

    return doc


# ============================================================
# GENEROWANIE OFERTY DOCX
# ============================================================

def generate_offer_docx(data, selected_products, template_data):
    """Generuj ofertę DOCX - multi-file (WolfTax)"""
    print(f"[DOCX] Generuję ofertę WolfTax z {len(template_data['files'])} plików")

    form_data = data.get('formData', data)
    product_custom_fields = data.get('productCustomFields', {})

    template_folder = os.path.join(TEMPLATES_DIR, template_data['folder'])
    files = sorted(template_data['files'], key=lambda x: x['order'])
    injection_point = template_data.get('injection_point', {})

    # Lista wszystkich dokumentów do połączenia
    docs_to_merge = []

    # Przetwórz wszystkie pliki
    for file_info in files:
        file_path = os.path.join(template_folder, file_info['file'])

        if not os.path.exists(file_path):
            continue

        doc = Document(file_path)
        doc = replace_placeholders(doc, form_data)

        # Spis treści
        if file_info.get('is_toc') and len(selected_products) > 0:
            toc_config = template_data.get('toc', {})
            start_page = toc_config.get('start_page', 5)
            toc_text = generate_table_of_contents(selected_products, product_custom_fields, start_page)
            doc = inject_toc_into_doc(doc, toc_text)
            print(f"[DOCX] Spis treści dodany")

        docs_to_merge.append(doc)

        # Injection point - produkty
        if (injection_point.get('type') == 'between_files' and
            file_info['file'] == injection_point.get('after')):

            # Wstaw produkty
            print(f"[DOCX] Wstawiam {len(selected_products)} produktów")
            for product_id in selected_products:
                product_path = os.path.join(PRODUKTY_DIR, f'{product_id}.docx')
                if os.path.exists(product_path):
                    product_doc = Document(product_path)

                    if product_id in product_custom_fields:
                        custom_data = product_custom_fields[product_id]
                        product_doc = replace_placeholders(product_doc, custom_data)

                    docs_to_merge.append(product_doc)

    # Połącz wszystkie dokumenty
    merged_doc = merge_documents(docs_to_merge)

    # Zapisz
    client_name = form_data.get('NazwaFirmyKlienta') or form_data.get('klient') or 'Klient'
    # Sanitize filename
    client_name = "".join(c for c in client_name if c.isalnum() or c in (' ', '-', '_')).strip()
    output_filename = f"Oferta_WolfTax_{client_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = os.path.join(GENERATED_OFFERS_DIR, output_filename)

    merged_doc.save(output_path)
    print(f"[DOCX] ✓ Zapisano: {output_filename}")

    return output_path, output_filename


# ============================================================
# API ROUTES
# ============================================================

@app.route('/')
def index():
    """Główna strona"""
    response = make_response(render_template('index.html'))
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    return response


@app.route('/api/templates')
def get_templates():
    """Lista szablonów"""
    templates_path = os.path.join(TEMPLATES_DIR, 'templates.json')
    with open(templates_path, 'r', encoding='utf-8') as f:
        return jsonify(json.load(f))


@app.route('/api/products')
def get_products():
    """Lista produktów"""
    products = []
    if os.path.exists(PRODUKTY_DIR):
        for filename in sorted(os.listdir(PRODUKTY_DIR)):
            if filename.endswith('.docx') and not filename.startswith('~$'):
                product_id = filename.replace('.docx', '')
                products.append({
                    'id': product_id,
                    'name': f'Produkt {product_id}',
                    'filename': filename
                })
    return jsonify(products)


@app.route('/api/save-offer', methods=['POST'])
def save_offer():
    """Zapisz ofertę do JSON"""
    data = request.json
    offer_name = data.get('offer_name', f"oferta_{datetime.now().strftime('%Y%m%d_%H%M%S')}")

    offer_data = {k: v for k, v in data.items() if k != 'offer_name'}

    filename = f"{offer_name}.json"
    filepath = os.path.join(SAVED_OFFERS_DIR, filename)

    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(offer_data, f, ensure_ascii=False, indent=2)

    return jsonify({'success': True, 'filename': filename})


@app.route('/api/load-offer/<filename>')
def load_offer(filename):
    """Wczytaj zapisaną ofertę"""
    filepath = os.path.join(SAVED_OFFERS_DIR, filename)

    if not os.path.exists(filepath):
        return jsonify({'error': 'Plik nie istnieje'}), 404

    with open(filepath, 'r', encoding='utf-8') as f:
        return jsonify(json.load(f))


@app.route('/api/saved-offers')
def get_saved_offers():
    """Lista zapisanych ofert"""
    offers = []
    if os.path.exists(SAVED_OFFERS_DIR):
        for filename in os.listdir(SAVED_OFFERS_DIR):
            if filename.endswith('.json'):
                filepath = os.path.join(SAVED_OFFERS_DIR, filename)
                stat = os.stat(filepath)
                offers.append({
                    'filename': filename,
                    'name': filename.replace('.json', ''),
                    'modified': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
                })

    offers.sort(key=lambda x: x['modified'], reverse=True)
    return jsonify(offers)


@app.route('/api/generate-offer', methods=['POST'])
def generate_offer():
    """Generuj DOCX"""
    import time
    start_time = time.time()

    data = request.json
    form_data = data.get('formData', {})
    selected_products = data.get('selectedProducts', [])
    template_data = data.get('templateData')

    try:
        send_progress("⚙️ Generowanie DOCX...", 10)

        output_path, output_filename = generate_offer_docx(data, selected_products, template_data)

        elapsed = time.time() - start_time
        send_progress(f"✅ Gotowe! ({elapsed:.1f}s)", 100)

        time.sleep(0.3)
        send_progress("", 0)

        return jsonify({
            'success': True,
            'filename': output_filename,
            'download_url': f'/api/download-offer/{output_filename}',
            'generation_time': f"{elapsed:.2f}s"
        })
    except Exception as e:
        send_progress(f"❌ Błąd: {str(e)}", 0)
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/download-offer/<filename>')
def download_offer(filename):
    """Pobierz DOCX"""
    filepath = os.path.join(GENERATED_OFFERS_DIR, filename)

    if not os.path.exists(filepath):
        return jsonify({'error': 'Plik nie istnieje'}), 404

    return send_file(filepath, as_attachment=True, download_name=filename)


@app.route('/api/preview-full-offer', methods=['POST'])
def preview_full_offer():
    """Generuj podgląd JPG"""
    data = request.json
    template_data = data.get('templateData')
    form_data = data.get('formData', {})
    selected_products = data.get('selectedProducts', [])
    product_custom_fields = data.get('productCustomFields', {})

    print(f"[PREVIEW] Template: {template_data['id']}, Produkty: {selected_products}")

    send_progress("Generuję podgląd...", 5)

    pages_metadata = []
    page_counter = 0

    # WolfTax multi-file
    template_folder = os.path.join(TEMPLATES_DIR, template_data['folder'])
    files = sorted(template_data['files'], key=lambda x: x['order'])
    injection_point = template_data.get('injection_point', {})

    # Przetwórz każdy plik
    for file_info in files:
        file_name = file_info['file']
        file_path = os.path.join(template_folder, file_name)

        if not os.path.exists(file_path):
            continue

        print(f"[PREVIEW] Przetwarzam: {file_name}")
        send_progress(f"📄 {file_info.get('name', file_name)}...", 10 + page_counter * 2)

        # Załaduj i wypełnij placeholders
        doc = Document(file_path)
        doc = replace_placeholders(doc, form_data)

        # Spis treści
        if file_info.get('is_toc') and len(selected_products) > 0:
            toc_config = template_data.get('toc', {})
            start_page = toc_config.get('start_page', 5)
            toc_text = generate_table_of_contents(selected_products, product_custom_fields, start_page)
            doc = inject_toc_into_doc(doc, toc_text)

        # Konwertuj na JPG
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx', dir=OUT_JPG_DIR) as temp_file:
            doc.save(temp_file.name)
            temp_path = temp_file.name

        file_images = convert_docx_to_images(temp_path, use_cache=False)

        try:
            os.unlink(temp_path)
        except:
            pass

        # Wyślij strony
        for idx, img_data in enumerate(file_images):
            page_counter += 1
            page_data = {
                'type': 'template',
                'number': page_counter,
                'image': img_data,
                'has_image': True,
                'page_index': idx,
                'status': 'ready',
                'source_file': file_name
            }
            pages_metadata.append(page_data)
            send_page_ready(page_data)

        # Injection point - produkty
        if (injection_point.get('type') == 'between_files' and
            file_info['file'] == injection_point.get('after')):

            print(f"[PREVIEW] Injection point - wstawiam {len(selected_products)} produktów")
            send_progress("Dodaję produkty...", 50)

            for product_id in selected_products:
                product_path = os.path.join(PRODUKTY_DIR, f'{product_id}.docx')

                if not os.path.exists(product_path):
                    continue

                # Custom fields
                path_to_convert = product_path
                temp_file = None

                if product_id in product_custom_fields and product_custom_fields[product_id]:
                    custom_data = product_custom_fields[product_id]
                    product_doc = Document(product_path)
                    product_doc = replace_placeholders(product_doc, custom_data)

                    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx', dir=OUT_JPG_DIR)
                    product_doc.save(temp_file.name)
                    temp_file.close()
                    path_to_convert = temp_file.name

                # Konwertuj
                use_cache = (temp_file is None)
                product_images = convert_docx_to_images(path_to_convert, use_cache=use_cache)

                if temp_file:
                    try:
                        os.unlink(temp_file.name)
                    except:
                        pass

                # Wyślij strony produktu
                for idx, img_data in enumerate(product_images):
                    page_counter += 1
                    page_data = {
                        'type': 'product',
                        'number': page_counter,
                        'product_id': product_id,
                        'image': img_data,
                        'has_image': True,
                        'page_index': idx,
                        'status': 'ready'
                    }
                    pages_metadata.append(page_data)
                    send_page_ready(page_data)

    send_progress("✅ Gotowe!", 100)

    import time
    time.sleep(0.3)
    send_progress("", 0)

    # Usuń obrazy z metadanych (są już wysłane przez WebSocket)
    metadata_without_images = []
    for meta in pages_metadata:
        meta_copy = {k: v for k, v in meta.items() if k != 'image'}
        meta_copy['has_image'] = meta.get('status') == 'ready'
        metadata_without_images.append(meta_copy)

    return jsonify({
        'success': True,
        'total_pages': len(pages_metadata),
        'pages_metadata': metadata_without_images
    })


# WebSocket
@socketio.on('connect')
def handle_connect():
    print(f'[WebSocket] ✅ Połączony: {request.sid}')

@socketio.on('disconnect')
def handle_disconnect():
    print(f'[WebSocket] ❌ Rozłączony: {request.sid}')


# ============================================================
# STARTUP
# ============================================================

print("\n" + "="*80)
print("🚀 ZOPTYMALIZOWANY GENERATOR OFERT")
print("="*80)
print(f"LibreOffice: {find_libreoffice() or 'NIE ZNALEZIONO'}")
print(f"PyMuPDF: {'✓ TAK' if HAS_PYMUPDF else '✗ NIE (używam pdf2image)'}")
print("="*80)

# Uruchom pre-rendering w tle
preload_async()

if __name__ == '__main__':
    socketio.run(app, debug=True, host='0.0.0.0', port=40207, allow_unsafe_werkzeug=True)
